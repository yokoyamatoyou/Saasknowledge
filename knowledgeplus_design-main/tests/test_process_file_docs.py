import io
import sys
from pathlib import Path

sys.path.insert(1, str(Path(__file__).resolve().parents[1]))

from shared.file_processor import FileProcessor  # noqa: E402
from shared.kb_builder import KnowledgeBuilder  # noqa: E402


def _create_text_pdf():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(40, 10, "hello pdf")
    data = pdf.output(dest="S").encode("latin1")
    buf = io.BytesIO(data)
    buf.name = "text.pdf"
    return buf


def _create_image_pdf(tmp_path):
    from fpdf import FPDF
    from PIL import Image

    img_path = tmp_path / "img.png"
    Image.new("RGB", (10, 10), "red").save(img_path)
    pdf = FPDF()
    pdf.add_page()
    pdf.image(str(img_path), x=10, y=10, w=10)
    data = pdf.output(dest="S").encode("latin1")
    buf = io.BytesIO(data)
    buf.name = "scan.pdf"
    return buf


def _create_text_docx(tmp_path):
    import docx

    path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("hello docx")
    doc.save(path)
    return path


def _create_image_docx(tmp_path):
    import docx
    from PIL import Image

    path = tmp_path / "img.docx"
    doc = docx.Document()
    img_path = tmp_path / "img.png"
    Image.new("RGB", (10, 10), "blue").save(img_path)
    doc.add_picture(str(img_path))
    doc.save(path)
    return path


def test_process_file_text_pdf():
    buf = _create_text_pdf()
    builder = KnowledgeBuilder(FileProcessor(), lambda: None, lambda *_: None)
    result = FileProcessor.process_file(buf, builder=builder)
    assert result["type"] == "document"
    assert "hello pdf" in result["text"]
    assert result["images"]


def test_process_file_image_pdf(tmp_path):
    buf = _create_image_pdf(tmp_path)
    builder = KnowledgeBuilder(FileProcessor(), lambda: None, lambda *_: None)
    result = FileProcessor.process_file(buf, builder=builder)
    assert result["type"] == "image"
    assert result["metadata"]["file_type"] == "image_file"
    assert result["image_base64"]


def test_process_file_text_docx(tmp_path):
    path = _create_text_docx(tmp_path)
    with open(path, "rb") as f:
        builder = KnowledgeBuilder(FileProcessor(), lambda: None, lambda *_: None)
        result = FileProcessor.process_file(f, builder=builder)
    assert result["type"] == "document"
    assert "hello docx" in result["text"]


def test_process_file_image_docx(tmp_path):
    path = _create_image_docx(tmp_path)
    with open(path, "rb") as f:
        builder = KnowledgeBuilder(FileProcessor(), lambda: None, lambda *_: None)
        result = FileProcessor.process_file(f, builder=builder)
    assert result["type"] == "image"
    assert result["metadata"]["file_type"] == "image_file"
    assert result["image_base64"]
